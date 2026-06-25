"""Ingestion pipeline (LangChain + hybrid Qdrant).

    load (robust, encoding-aware) -> Documents (+metadata)
      -> header-aware chunking (Markdown) / recursive chunking (other)
      -> dense BGE-M3 + sparse BM25 -> QdrantVectorStore (HYBRID, recreated)

Run inside the backend container:
    docker compose run --rm backend python -m ingestion.ingest
Options:
    --data DIR   source directory (default: data)

The collection is always rebuilt from scratch (hybrid named-vector schema), so
deletions in data/ are reflected. Chunks are de-duplicated by content hash.
"""
import argparse
import hashlib
import re
import sys
from pathlib import Path
from typing import Iterable, List

import ftfy
from langchain_core.documents import Document
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)

from app.config import settings
from app.qdrant_store import build_collection

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".csv"}

_RECURSIVE = RecursiveCharacterTextSplitter(
    chunk_size=900,
    chunk_overlap=150,
    separators=["\n## ", "\n### ", "\n\n", "\n", ". ", " ", ""],
)
_MD_HEADERS = MarkdownHeaderTextSplitter(
    headers_to_split_on=[("#", "h1"), ("##", "h2"), ("###", "h3")],
    strip_headers=True,
)


# ---------- load ----------
def read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def read_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def raw_corruption_ratio(path: Path) -> float:
    raw = path.read_bytes()
    if not raw:
        return 1.0
    return raw.count(0x3F) / len(raw)


def read_text(path: Path) -> str:
    raw = path.read_bytes()
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16")
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    from charset_normalizer import from_bytes
    best = from_bytes(raw).best()
    if best is not None:
        enc = (best.encoding or "").lower()
        if not enc.startswith("utf_16") and not enc.startswith("utf-16"):
            return str(best)
    return raw.decode("cp1251", errors="ignore")


def read_csv(path: Path) -> str:
    text = read_text(path)
    import csv
    import io
    delim = ";" if text[:2000].count(";") >= text[:2000].count(",") else ","
    rows = csv.reader(io.StringIO(text), delimiter=delim)
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def load_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".csv":
        return read_csv(path)
    if ext in (".txt", ".md"):
        return read_text(path)
    return ""


# ---------- front-matter (source_url + optional metadata overrides) ----------
# Lets a .md/.txt file declare its origin URL (and optionally override title,
# lang, etc.) in a leading YAML-style block, so RAG answers can cite real links.
#   ---
#   source_url: https://www.alatoo.edu.kg/admission
#   title: Поступление
#   ---
_FM_OVERRIDE_KEYS = {
    "source_url", "title", "source", "lang", "doc_type", "faculty", "program",
}


def parse_front_matter(text: str) -> tuple[dict, str]:
    """Parse a leading '---' fenced block of simple ``key: value`` lines.

    Returns ``(meta, body)``. Only keys in ``_FM_OVERRIDE_KEYS`` are kept. If
    there is no opening fence, no closing fence, or no recognised key, the text
    is returned untouched (so a leading Markdown horizontal rule is not eaten).
    """
    stripped = text.lstrip("﻿").lstrip()
    if not stripped.startswith("---"):
        return {}, text
    lines = stripped.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, text
    meta: dict = {}
    end = None
    for i in range(1, len(lines)):
        if lines[i].strip() in ("---", "..."):
            end = i
            break
        m = re.match(r"\s*([A-Za-z_][\w-]*)\s*:\s*(.*?)\s*$", lines[i])
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip().strip('"').strip("'")
            if key in _FM_OVERRIDE_KEYS and val:
                meta[key] = val
    if end is None or not meta:
        return {}, text
    body = "\n".join(lines[end + 1:])
    return meta, body


# ---------- clean ----------
def clean_text(text: str) -> str:
    text = ftfy.fix_text(text)
    text = text.replace("­", "")      # soft hyphens
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def guess_lang(text: str) -> str:
    sample = text[:2000]
    cyr = len(re.findall(r"[а-яёңүөъ]", sample, flags=re.IGNORECASE))
    lat = len(re.findall(r"[a-z]", sample, flags=re.IGNORECASE))
    if re.search(r"[ңүөъ]", sample, flags=re.IGNORECASE):
        return "ky"
    if cyr > lat:
        return "ru"
    if lat > 0:
        return "en"
    return "ru"


def _broken_ratio(text: str) -> float:
    if not text:
        return 1.0
    bad = text.count("?") + text.count("�")
    return bad / len(text)


# ---------- chunk -> Documents ----------
def _faculty_from(headers: dict) -> str:
    for v in headers.values():
        if v and "факультет" in v.lower():
            return v.strip()
    return ""


# Per-section override: a multi-topic file (e.g. the compiled university
# handbook) can point each section at its own canonical page by placing an
# HTML comment right under the heading:  <!-- source_url: https://... -->
# Falls back to the file-level front-matter source_url when absent.
_SECTION_URL_RE = re.compile(r"<!--\s*source_url:\s*(\S+?)\s*-->")


def docs_from_markdown(text: str, base_meta: dict) -> List[Document]:
    out: List[Document] = []
    file_url = base_meta.get("source_url", "")
    by_h2: dict = {}                       # H2-level URL, inherited by its H3 children
    for sec in _MD_HEADERS.split_text(text):
        headers = {k: v for k, v in sec.metadata.items() if k in ("h1", "h2", "h3")}
        trail = " > ".join(headers[k] for k in ("h1", "h2", "h3") if headers.get(k))
        section = headers.get("h3") or headers.get("h2") or headers.get("h1") or ""
        faculty = _faculty_from(headers)
        h2 = headers.get("h2", "")
        has_h3 = bool(headers.get("h3"))
        sec_text = sec.page_content
        m = _SECTION_URL_RE.search(sec_text)
        if m:
            sec_url = m.group(1)
            sec_text = (sec_text[:m.start()] + sec_text[m.end():])
            if h2 and not has_h3:          # marker on an H2 intro -> inherited by its H3s
                by_h2[h2] = sec_url
        else:
            sec_url = by_h2.get(h2, file_url)
        for piece in _RECURSIVE.split_text(sec_text):
            content = (trail + "\n" + piece).strip() if trail else piece.strip()
            md = dict(base_meta)
            md.update({"section": section, "faculty": faculty,
                       "source_url": sec_url,
                       "title": section or base_meta.get("title", "")})
            out.append(Document(page_content=content, metadata=md))
    return out


def docs_from_plain(text: str, base_meta: dict) -> List[Document]:
    return [
        Document(page_content=p.strip(), metadata=dict(base_meta))
        for p in _RECURSIVE.split_text(text) if p.strip()
    ]


def iter_files(data_dir: Path) -> Iterable[Path]:
    for p in sorted(data_dir.rglob("*")):
        if p.is_file() and p.suffix.lower() in SUPPORTED:
            yield p


def build_documents(data_dir: Path) -> List[Document]:
    files = list(iter_files(data_dir))
    if not files:
        print("[!] No supported files in " + str(data_dir), file=sys.stderr)
        sys.exit(1)

    print("[i] Found " + str(len(files)) + " files.")
    seen = set()
    documents: List[Document] = []

    for path in files:
        ext = path.suffix.lower()
        if ext in (".txt", ".md", ".csv") and raw_corruption_ratio(path) >= 0.2:
            print("  - SKIP (corrupted encoding): " + path.name)
            continue
        raw = load_file(path)
        fm = {}
        if ext in (".md", ".txt"):
            fm, raw = parse_front_matter(raw)
        text = clean_text(raw)
        if not text:
            print("  - skip (empty): " + path.name)
            continue
        if _broken_ratio(text) >= 0.2:
            print("  - SKIP (corrupted text): " + path.name)
            continue

        lang = guess_lang(text)
        base_meta = {
            "source": path.name,
            "title": path.stem,
            "doc_type": ext.lstrip("."),
            "lang": lang,
            "source_url": "",
            "updated_at": int(path.stat().st_mtime),
        }
        for k, v in fm.items():          # front-matter overrides (source_url, title, ...)
            base_meta[k] = v
        docs = docs_from_markdown(text, base_meta) if ext == ".md" else docs_from_plain(text, base_meta)

        kept = 0
        for d in docs:
            h = hashlib.sha256(d.page_content.encode("utf-8")).hexdigest()
            if h in seen:
                continue
            seen.add(h)
            documents.append(d)
            kept += 1
        print("  + " + path.name + ": " + str(kept) + " chunks (" + lang + ")")

    return documents


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", default="data")
    ap.add_argument("--recreate", action="store_true",
                    help="(kept for compatibility; collection is always rebuilt)")
    args = ap.parse_args()

    data_dir = Path(args.data)
    if not data_dir.exists():
        print("[!] Data dir '" + str(data_dir) + "' not found.", file=sys.stderr)
        sys.exit(1)

    documents = build_documents(data_dir)
    if not documents:
        print("[!] No documents to index.", file=sys.stderr)
        sys.exit(1)

    print("[i] Building hybrid collection '" + settings.qdrant_collection +
          "' (dense=" + settings.embed_model + " + sparse=" + settings.sparse_model +
          ") from " + str(len(documents)) + " chunks...")
    build_collection(documents)
    print("[OK] Done. Indexed " + str(len(documents)) + " chunks (hybrid dense+sparse).")


if __name__ == "__main__":
    main()
